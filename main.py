import streamlit as st
import pandas as pd
from fpdf import FPDF
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.section import WD_ORIENT
from docx.shared import Inches
import io
import os

ADMIN_PASSWORD = "12"
DB_FILE = "database.csv" # SAVES DATA SO ALL PHONES SEE IT

st.set_page_config(page_title="MADAM RELINDIS DATABASE SYSTEM", layout="wide")

# ===== YOUR FULL DESIGN - COLORS + BORDERS + BOLD =====
st.markdown("""
<style>
html, body, div, input, button, p, label, h1, h2, h3, h4, h5, h6 {font-weight: 900!important;}
.animated-title {text-align: center; font-size: 38px; font-weight: 900; background: linear-gradient(90deg, #FF0000, #FF7F00, #FFFF00, #00FF00, #0000FF, #4B0082, #8B00FF); background-size: 400% 400%; -webkit-background-clip: text; -webkit-text-fill-color: transparent; animation: rainbow 3s linear infinite; padding: 20px; border: 4px solid #1E3A8A; border-radius: 12px; background-color: #F0F8FF;}
@keyframes rainbow {0%{background-position:0% 50%}100%{background-position:100% 50%}}
.card {border: 3px solid #1E3A8A; border-radius: 10px; padding: 20px; margin-bottom: 20px; background-color: #FFFFFF;}
.stDataFrame {border: 3px solid #1E3A8A!important; border-radius: 8px;}
thead tr th {background-color: #1E3A8A!important; color: white!important; font-weight: 900!important; border: 2px solid white!important; text-align:center!important;}
tbody tr td {font-weight: 900!important; border: 1px solid #000!important;}
.stTextInput > div > div > input {border: 2px solid #1E3A8A!important;}
div[data-testid="stButton"] > button, div[data-testid="stDownloadButton"] > button {border: 3px solid black!important; height: 45px; border-radius: 8px; font-weight: 900!important;}
</style>
""", unsafe_allow_html=True)

# ===== LOAD/SAVE FUNCTIONS =====
def load_data():
    if os.path.exists(DB_FILE):
        return pd.read_csv(DB_FILE)
    else:
        return pd.DataFrame(columns=["Item", "Quantity", "Unit Price", "Total", "Purpose", "Statute"])

def save_data(df):
    df.to_csv(DB_FILE, index=False)

# ===== CLEAN DATAFRAME FUNCTION =====
def clean_df(df):
    if df.empty:
        return df
    df = df.copy()
    df.columns = df.columns.astype(str)
    if len(df.columns) >= 2:
        df = df.iloc[:, :-2]
    if len(df.columns) == 6:
        df.columns = ["Item", "Quantity", "Unit Price", "Total", "Purpose", "Statute"]
    return df

# ===== SESSION STATE =====
if "df" not in st.session_state:
    st.session_state.df = load_data() # LOAD FROM FILE
if "admin_mode" not in st.session_state:
    st.session_state.admin_mode = False
if "edit_mode" not in st.session_state:
    st.session_state.edit_mode = False
if "failed_files" not in st.session_state:
    st.session_state.failed_files = []

# ===== TITLE =====
st.markdown('<div class="card"><h1 class="animated-title">MADAM RELINDIS DATABASE SYSTEM</h1></div>', unsafe_allow_html=True)

# ===== ALL 6 BUTTONS WITH ACTIONS =====
st.markdown('<div class="card">', unsafe_allow_html=True)
col1, col2, col3 = st.columns(3)

with col1:
    if st.button("Import Excel Files", use_container_width=True):
        st.session_state.admin_mode = True
    if st.button("Edit", use_container_width=True): # NEW EDIT BUTTON
        st.session_state.edit_mode = True

with col2:
    csv = clean_df(st.session_state.df).to_csv(index=False).encode('utf-8')
    st.download_button("Export Full Database", csv, "Madam_Relindis_Database.csv", "text/csv", use_container_width=True)

with col3:
    if st.button("Export Word Report", use_container_width=True):
        st.session_state.export_word = True

col4, col5 = st.columns(2)
with col4:
    if st.button("Export PDF Report", use_container_width=True):
        st.session_state.export_pdf = True

with col5:
    if st.button("View Failed Files", use_container_width=True):
        st.session_state.view_failed = True
st.markdown('</div>', unsafe_allow_html=True)

# BUTTON COLORS
st.markdown("""
<style>
div.stButton > button {background: #FF6B35!important; color: white!important;}
div.stButton > button:nth-of-type(2) {background: #00BFFF!important; color: white!important;} /* EDIT BUTTON BLUE */
div.stButton > button:nth-of-type(3) {background: #2196F3!important; color: white!important;}
div.stButton > button:nth-of-type(4) {background: #E91E63!important; color: white!important;}
div.stButton > button:nth-of-type(5) {background: #FFC107!important; color: black!important;}
div.stButton > button:nth-of-type(6) {background: #9C27B0!important; color: white!important;}
div.stDownloadButton > button {background: #4CAF50!important; color: white!important;}
</style>
""", unsafe_allow_html=True)

# ===== WORD EXPORT =====
if "export_word" in st.session_state and st.session_state.export_word:
    df_export = clean_df(st.session_state.df)
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69); section.page_height = Inches(8.27)
    section.left_margin = Inches(0.3); section.right_margin = Inches(0.3)
    title = doc.add_heading('MADAM RELINDIS DATABASE REPORT', 0)
    title.alignment = 1
    for run in title.runs: run.font.underline = True; run.font.bold = True
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=len(df_export.columns))
    table.style = 'Table Grid'; table.autofit = False
    total_width_emu = int(11 * 914400)
    col_width_emu = int(total_width_emu / len(df_export.columns))
    for i in range(len(df_export.columns)): table.columns[i].width = col_width_emu
    for i, col in enumerate(df_export.columns):
        cell = table.cell(0, i); cell.text = str(col); cell.paragraphs[0].alignment = 1
        for run in cell.paragraphs[0].runs: run.font.bold = True; run.font.size = Pt(8); run.font.color.rgb = RGBColor(255, 255, 255)
        shading_elm = parse_xml(r'<w:shd {} w:fill="1E3A8A"/>'.format(nsdecls('w'))); cell._tc.get_or_add_tcPr().append(shading_elm)
    for _, row in df_export.iterrows():
        row_cells = table.add_row().cells
        for i, item in enumerate(row): cell = row_cells[i]; cell.text = str(item); cell.paragraphs[0].alignment = 1; run = cell.paragraphs[0].add_run(); run.font.size = Pt(8)
    buffer = io.BytesIO(); doc.save(buffer)
    st.download_button("Download Word File", buffer.getvalue(), "Report.docx")
    st.session_state.export_word = False

# ===== PDF EXPORT =====
if "export_pdf" in st.session_state and st.session_state.export_pdf:
    df_export = clean_df(st.session_state.df)
    pdf = FPDF(orientation='L', unit='mm', format='A4')
    pdf.add_page(); pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, "MADAM RELINDIS DATABASE REPORT", 0, 1, 'C'); pdf.ln(3)
    pdf.set_font("Arial", 'B', 7); col_width = 277 / len(df_export.columns)
    for col in df_export.columns: pdf.cell(col_width, 8, str(col)[:18], 1, 0, 'C')
    pdf.ln(); pdf.set_font("Arial", '', 7)
    for _, row in df_export.iterrows():
        for item in row: pdf.cell(col_width, 8, str(item)[:18], 1, 0, 'C')
        pdf.ln()
    buffer = io.BytesIO(); pdf.output(buffer)
    st.download_button("Download PDF File", buffer.getvalue(), "Report.pdf")
    st.session_state.export_pdf = False

# ===== VIEW FAILED FILES =====
if "view_failed" in st.session_state and st.session_state.view_failed:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("❌ Failed Files Log")
    if len(st.session_state.failed_files) == 0:
        st.success("No failed files. Everything is good!")
    else:
        for f in st.session_state.failed_files: st.error(f)
    if st.button("Close Failed Files"): st.session_state.view_failed = False; st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

# ===== ADMIN PANEL - IMPORT =====
if st.session_state.admin_mode:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("🔒 Admin Panel - Import")
    pwd = st.text_input("Enter Admin Password", type="password", key="admin_pwd")
    uploaded_file = st.file_uploader("Choose Excel File to Import", type=["xlsx", "xls"])
    col_a, col_b = st.columns(2)
    with col_a:
        if st.button("Confirm Import"):
            if pwd == ADMIN_PASSWORD:
                if uploaded_file:
                    try:
                        raw_df = pd.read_excel(uploaded_file)
                        if st.session_state.df.empty:
                            st.session_state.df = raw_df
                        else:
                            st.session_state.df = pd.concat([st.session_state.df, raw_df], ignore_index=True)
                        save_data(st.session_state.df) # SAVE TO FILE
                        st.success("1 file(s) imported successfully")
                        st.session_state.admin_mode = False; st.rerun()
                    except Exception as e:
                        st.session_state.failed_files.append(f"{uploaded_file.name}: {str(e)}")
                        st.error(f"Failed to import {uploaded_file.name}")
                else:
                    st.warning("Please select a file first")
            else:
                st.error("Wrong password")
    with col_b:
        if st.button("Lock Admin"): st.session_state.admin_mode = False; st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

# ===== EDIT PANEL =====
if st.session_state.edit_mode:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.subheader("✏️ Edit Data - Password Required")
    edit_pwd = st.text_input("Enter Edit Password", type="password", key="edit_pwd")

    if edit_pwd == ADMIN_PASSWORD:
        st.success("Access Granted. You can edit directly in the table below")
        edited_df = st.data_editor(st.session_state.df, num_rows="dynamic", use_container_width=True, key="editor")
        if st.button("Save Changes"):
            st.session_state.df = edited_df
            save_data(st.session_state.df) # SAVE TO FILE
            st.success("Changes Saved! All phones will now see this data")
            st.session_state.edit_mode = False
            st.rerun()
    elif edit_pwd:
        st.error("Wrong password")

    if st.button("Close Edit"):
        st.session_state.edit_mode = False
        st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

# ===== SEARCH + TABLE =====
st.markdown('<div class="card">', unsafe_allow_html=True)
search_term = st.text_input("🔍 Search", placeholder="Type here to search...")
df_display = clean_df(st.session_state.df)
if search_term and not df_display.empty:
    df_display = df_display[df_display.apply(lambda row: row.astype(str).str.contains(search_term, case=False, na=False).any(), axis=1)]
st.dataframe(df_display, use_container_width=True, height=400)
st.markdown('</div>', unsafe_allow_html=True)
